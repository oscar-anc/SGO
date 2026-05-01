"""
Helper Functions for Document Generation

Responsibilities:
  1. Transform formData (from all pages) into a clean Jinja2 context dict.
  2. Build complex tables (policy tables, financing tables) as python-docx
     subdocuments and inject them into the context as {{subdocPolicyTables}}
     and {{subdocFinancingTables}}.
  3. Render template.docx with docxtpl and save to the user-chosen path.

Style constants here must stay in sync with build_template.js.
If you change a color or width in build_template.js, update the matching
constant below and regenerate the template with: node build_template.js
"""

from datetime import date

from docxtpl import DocxTemplate
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

from PySide6.QtWidgets import QFileDialog
import sys as _sys


# ─────────────────────────────────────────────────────────────────────────────
# RESOURCE PATH
# Resolves file paths correctly both when running as a Python script and when
# bundled as a PyInstaller --onefile executable (sys._MEIPASS is set by the
# bootloader only in the frozen/exe context).
# ─────────────────────────────────────────────────────────────────────────────

def resource_path(relative_path: str) -> str:
    """Return absolute path to a resource, works for dev and PyInstaller."""
    import os as _os
    if getattr(_sys, 'frozen', False):
        base = _sys._MEIPASS
    else:
        base = _os.path.dirname(_os.path.abspath(_os.path.join(_os.getcwd(), __file__)))
    return _os.path.join(base, relative_path)


# ─────────────────────────────────────────────────────────────────────────────
# TABLE STYLE CONSTANTS
# Keep in sync with build_template.js CONSTANTS section.
# ─────────────────────────────────────────────────────────────────────────────

FONT_NAME = "Arial"
FONT_SIZE = Pt(10)

# Brand colors as (R, G, B) tuples
COLOR_HEADER_BG   = (0x00, 0x0F, 0x47)   # #000F47 — table header background
COLOR_HEADER_TEXT = (0xCD, 0xEC, 0xFF)   # #CDECFF — table header text
COLOR_BORDER      = (0x00, 0x0F, 0x47)   # #000F47 — table borders
COLOR_BLACK       = (0x00, 0x00, 0x00)   # #000000 — body text

# A4 content width in cm (21cm page - 2x2.5cm margins = 16cm)
CONTENT_WIDTH_CM = 16.0

# Cell padding in DXA — must match cellMargins in build_template.js
# 110 DXA top/bottom (≈1.9mm), 120 DXA left/right (≈2.1mm)
CELL_MARGIN_TB  = 110   # top and bottom padding in DXA (twips)
CELL_MARGIN_LR  = 120   # left and right padding in DXA (twips)

# Column widths in cm — must sum to CONTENT_WIDTH_CM
# Policy table: RAMO | POLIZA | LIQUIDACION | PRIMA NETA | PRIMA TOTAL
COL_POLICY_CM = [3.52, 3.52, 3.17, 2.88, 2.89]

# Financing table 3-col: CUOTA N | FECHA VENCIMIENTO | IMPORTE
COL_FINANCING_CM = [3.87, 6.48, 5.65]

# Financing table 4-col (ambos mode): CUOTA | NRO.RECIBO | FECHA VENCIMIENTO | IMPORTE
# Total must sum to CONTENT_WIDTH_CM (16cm). Proportions: 2.0 + 4.0 + 5.0 + 5.0 = 16.0
COL_FINANCING_AMBOS_CM = [2.0, 4.0, 5.0, 5.0]

# Annex table: ANEXO | CONTENIDO
COL_ANNEX_CM = [2.12, 13.88]


# ─────────────────────────────────────────────────────────────────────────────
# LOW-LEVEL XML HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def _setRgbColor(cell, r, g, b):
    """Set solid RGB background on a table cell via XML (python-docx limitation)."""
    tc    = cell._tc
    tcPr  = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    shd      = OxmlElement('w:shd')
    hexColor = f'{r:02X}{g:02X}{b:02X}'
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hexColor)
    tcPr.append(shd)


def _setCellBorders(cell, top=True, bottom=True, left=True, right=True):
    """Set visible/hidden borders on a table cell via XML."""
    tc      = cell._tc
    tcPr    = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    tcBords   = OxmlElement('w:tcBorders')
    borderHex = f'{COLOR_BORDER[0]:02X}{COLOR_BORDER[1]:02X}{COLOR_BORDER[2]:02X}'

    def _b(name, visible):
        el = OxmlElement(f'w:{name}')
        if visible:
            el.set(qn('w:val'),   'single')
            el.set(qn('w:sz'),    '4')
            el.set(qn('w:color'), borderHex)
        else:
            el.set(qn('w:val'), 'nil')
        return el

    tcBords.append(_b('top',    top))
    tcBords.append(_b('left',   left))
    tcBords.append(_b('bottom', bottom))
    tcBords.append(_b('right',  right))
    tcPr.append(tcBords)


def _setCellMargins(cell, top=None, bottom=None, left=None, right=None):
    """
    Set internal cell padding (tcMar) via XML.
    Values in DXA (twips: 1/20 of a point, same unit as DXA in docx-js).
    Call after _setCellBorders so tcPr already exists.
    Defaults to CELL_MARGIN_TB (top/bottom) and CELL_MARGIN_LR (left/right)
    from the module-level constants, keeping them in sync with build_template.js.
    """
    tc   = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    tcMar = OxmlElement('w:tcMar')
    # Use provided values or fall back to module constants
    vals = {
        'top':    top    if top    is not None else CELL_MARGIN_TB,
        'bottom': bottom if bottom is not None else CELL_MARGIN_TB,
        'left':   left   if left   is not None else CELL_MARGIN_LR,
        'right':  right  if right  is not None else CELL_MARGIN_LR,
    }
    for side, dxa in vals.items():
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'),    str(dxa))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def _setTableWidth(table, width_cm=None, center=False):
    """
    Force table to a specific width.
    Defaults to CONTENT_WIDTH_CM (full page content width).
    Pass width_cm explicitly for narrower tables (e.g. financing table at 12cm).
    Pass center=True to horizontally center the table on the page.
    """
    w_cm  = width_cm if width_cm is not None else CONTENT_WIDTH_CM
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW  = OxmlElement('w:tblW')
    tblW.set(qn('w:w'),    str(int(w_cm * 567)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    # Optional horizontal centering via w:jc
    if center:
        tblJc = OxmlElement('w:jc')
        tblJc.set(qn('w:val'), 'center')
        tblPr.append(tblJc)


def _setColumnWidths(table, widths_cm):
    """Set exact column widths on table via XML tblGrid."""
    tbl     = table._tbl
    tblGrid = OxmlElement('w:tblGrid')
    for w in widths_cm:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(int(w * 567)))
        tblGrid.append(col)
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is not None:
        tblPr.addnext(tblGrid)
    else:
        tbl.insert(0, tblGrid)


# ─────────────────────────────────────────────────────────────────────────────
# CELL STYLER
# ─────────────────────────────────────────────────────────────────────────────

def _styleCell(cell, text, width_cm,
               bgColor=None, textColor=None,
               bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER,
               borders=None, allCaps=True):
    """
    Apply all visual styles to an existing table cell and write its text.

    Args:
        cell:       docx TableCell
        text:       string content
        width_cm:   column width in cm
        bgColor:    (R,G,B) tuple — None defaults to white
        textColor:  (R,G,B) tuple — None defaults to black
        bold:       bold font
        alignment:  WD_ALIGN_PARAGRAPH constant
        borders:    dict {top, bottom, left, right: bool} — None = all visible
        allCaps:    uppercase the text
    """
    # Background
    _setRgbColor(cell, *(bgColor or (0xFF, 0xFF, 0xFF)))

    # Width
    cell.width = Cm(width_cm)

    # Borders
    if borders is not None:
        _setCellBorders(cell, **borders)
    else:
        _setCellBorders(cell)

    # Cell padding — uses module constants (CELL_MARGIN_TB / CELL_MARGIN_LR)
    # which are kept in sync with cellMargins in build_template.js
    _setCellMargins(cell)

    # Vertical alignment
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Clear existing paragraph content
    para = cell.paragraphs[0]
    for child in para._p:
        para._p.remove(child)

    para.alignment = alignment
    para.paragraph_format.space_after  = Pt(0)
    para.paragraph_format.space_before = Pt(0)

    run = para.add_run(text.upper() if allCaps else text)
    run.font.name      = FONT_NAME
    run.font.size      = FONT_SIZE
    run.font.bold      = bold
    run.font.color.rgb = RGBColor(*(textColor or COLOR_BLACK))


def _setHeaderRowInnerBorders(row, num_cols):
    """
    Set inner vertical borders to white on all cells of a header row.
    In Word, adjacent cells share a border — the right border of cell N
    and the left border of cell N+1 are resolved by conflict rules.
    Setting right=white on every non-last cell and left=white on every
    non-first cell creates a clean visual separator between header columns.
    """
    WHITE_HEX = 'FFFFFF'
    for i, cell in enumerate(row.cells[:num_cols]):
        tc   = cell._tc
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            tc.insert(0, tcPr)
        tcBords = tcPr.find(qn('w:tcBorders'))
        if tcBords is None:
            tcBords = OxmlElement('w:tcBorders')
            tcPr.append(tcBords)

        def _white_border(name):
            existing = tcBords.find(qn(f'w:{name}'))
            if existing is not None:
                tcBords.remove(existing)
            el = OxmlElement(f'w:{name}')
            el.set(qn('w:val'),   'single')
            el.set(qn('w:sz'),    '4')
            el.set(qn('w:color'), WHITE_HEX)
            tcBords.append(el)

        # Right border white on all but last cell
        if i < num_cols - 1:
            _white_border('right')
        # Left border white on all but first cell
        if i > 0:
            _white_border('left')


def _addSpacerPara(doc):
    """Add an empty paragraph spacer between tables in a subdocument."""
    p   = doc.add_paragraph("")
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("")
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE


# ─────────────────────────────────────────────────────────────────────────────
# POLICY TABLE BUILDER
# ─────────────────────────────────────────────────────────────────────────────

def _sumPolicyColumn(polizas, key):
    """
    Sum numeric values from a list of policy dicts for a given key.
    Values may be formatted strings like '1,234.56' or '1234.56'.
    Returns formatted string with 2 decimal places, or '0.00' on failure.
    """
    import re
    total = 0.0
    for p in polizas:
        raw = p.get(key, '') or ''
        # Strip non-numeric except decimal point and minus
        cleaned = re.sub('[^0-9.-]', '', raw.replace(',', ''))
        try:
            total += float(cleaned)
        except (ValueError, TypeError):
            pass
    # Format with thousands separator and 2 decimals
    return f'{total:,.2f}'


def _buildOnePolicyTable(doc, polizas, currency, insuredName=None, liq_header='LIQUIDACIÓN'):
    """
    Add one policy table to a python-docx Document (or subdocument).

    insuredName: when provided, adds a merged name row at the top.
    This is only used when multipleClients is active.
    """
    # +1 total row when more than 1 policy; polizas guaranteed non-empty here
    hasTotalRow = len(polizas) > 1
    numRows = (1 if insuredName else 0) + 1 + max(len(polizas), 1) + (1 if hasTotalRow else 0)
    table   = doc.add_table(rows=numRows, cols=5)
    _setTableWidth(table)
    _setColumnWidths(table, COL_POLICY_CM)

    rowIdx = 0

    # Merged insured name row
    if insuredName:
        row = table.rows[rowIdx]
        row.cells[0].merge(row.cells[4])
        _styleCell(
            row.cells[0], insuredName, CONTENT_WIDTH_CM,
            bgColor=COLOR_HEADER_BG, textColor=COLOR_HEADER_TEXT, bold=True
        )
        rowIdx += 1

    # Column header row
    headers = [
        "RAMO", "PÓLIZA", liq_header,
        f"PRIMA NETA {currency}", f"PRIMA TOTAL {currency}"
    ]
    headerRow = table.rows[rowIdx]
    for col, (text, width) in enumerate(zip(headers, COL_POLICY_CM)):
        _styleCell(
            headerRow.cells[col], text, width,
            bgColor=COLOR_HEADER_BG, textColor=COLOR_HEADER_TEXT, bold=True
        )
    _setHeaderRowInnerBorders(headerRow, len(headers))
    rowIdx += 1

    # Data rows — Prima Neta and Prima Total are right-aligned per spec
    alignments = [
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.RIGHT,
        WD_ALIGN_PARAGRAPH.RIGHT,
    ]
    for policy in (polizas or [{}]):
        values  = [
            policy.get('POLIZA_RAMO',        ''),
            policy.get('POLIZA_NUMERO',      ''),
            policy.get('POLIZA_RECIBO',      ''),
            policy.get('POLIZA_PRIMA',       ''),
            policy.get('POLIZA_PRIMA_TOTAL', ''),
        ]
        dataRow = table.rows[rowIdx]
        for col, (val, width, align) in enumerate(
                zip(values, COL_POLICY_CM, alignments)):
            _styleCell(dataRow.cells[col], val, width, alignment=align)
        rowIdx += 1

    # Total row — only when more than 1 policy
    if len(polizas) > 1:
        totalRow = table.rows[rowIdx]
        # Col 1 — empty, top border only (no left/right/bottom)
        _styleCell(
            totalRow.cells[0], '', COL_POLICY_CM[0],
            borders={'top': True, 'bottom': False, 'left': False, 'right': False}
        )
        # Col 2 — empty, no bottom or left border
        _styleCell(
            totalRow.cells[1], '', COL_POLICY_CM[1],
            borders={'top': True, 'bottom': False, 'left': False, 'right': True}
        )
        # Col 3 — TOTAL label
        _styleCell(totalRow.cells[2], 'TOTAL', COL_POLICY_CM[2], bold=True)
        # Col 4 — Prima Neta sum (bold, right-aligned)
        _styleCell(
            totalRow.cells[3],
            _sumPolicyColumn(polizas, 'POLIZA_PRIMA'),
            COL_POLICY_CM[3],
            bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT
        )
        # Col 5 — Prima Total sum (bold, right-aligned)
        _styleCell(
            totalRow.cells[4],
            _sumPolicyColumn(polizas, 'POLIZA_PRIMA_TOTAL'),
            COL_POLICY_CM[4],
            bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT
        )


def _buildPolicySubdoc(docTemplate, formData, liq_header='LIQUIDACIÓN'):
    """
    Build all policy tables as a docxtpl subdocument.
    Replaces {{subdocPolicyTables}} in the template.
    """
    subdoc       = docTemplate.new_subdoc()
    currency     = formData.get('currency', '')
    multiInsured = formData.get('multipleClients', '') == 'Si'

    # Collect main insured policies
    mainPolicy = {
        'POLIZA_RAMO':        formData.get('POLIZA_RAMO',        ''),
        'POLIZA_NUMERO':      formData.get('POLIZA_NUMERO',      ''),
        'POLIZA_RECIBO':      formData.get('POLIZA_RECIBO',      ''),
        'POLIZA_PRIMA':       formData.get('POLIZA_PRIMA',       ''),
        'POLIZA_PRIMA_TOTAL': formData.get('POLIZA_PRIMA_TOTAL', ''),
    }
    allMainPolicies = [mainPolicy] + formData.get('polizas', [])
    mainInsuredName = formData.get('POLIZA_ASEGURADO', '') if multiInsured else None

    _buildOnePolicyTable(subdoc, allMainPolicies, currency, insuredName=mainInsuredName, liq_header=liq_header)

    # One table per additional insured
    for insured in formData.get('ASEGURADOS_ADICIONALES', []):
        _addSpacerPara(subdoc)
        insuredPolicy = {
            'POLIZA_RAMO':        insured.get('POLIZA_RAMO',        ''),
            'POLIZA_NUMERO':      insured.get('POLIZA_NUMERO',      ''),
            'POLIZA_RECIBO':      insured.get('POLIZA_RECIBO',      ''),
            'POLIZA_PRIMA':       insured.get('POLIZA_PRIMA',       ''),
            'POLIZA_PRIMA_TOTAL': insured.get('POLIZA_PRIMA_TOTAL', ''),
        }
        allPolicies = [insuredPolicy] + insured.get('polizas', [])
        _buildOnePolicyTable(
            subdoc, allPolicies, currency,
            insuredName=insured.get('POLIZA_ASEGURADO', ''),
            liq_header=liq_header
        )

    return subdoc


# ─────────────────────────────────────────────────────────────────────────────
# FINANCING TABLE BUILDER
# ─────────────────────────────────────────────────────────────────────────────

# ─────────────────────────────────────────────────────────────────────────────
# FINANCING TABLE BUILDER
# ─────────────────────────────────────────────────────────────────────────────

def _buildOneFinancingTable(doc, cuotas, total, currency,
                             titleRow=None, subtitleRow=None, modoRecibo='cuota'):
    """
    Add one financing table to a python-docx Document (or subdocument).

    titleRow:    optional first merged header (insured name or branch)
    subtitleRow: optional second merged header (branch within insured, Case 4 Individual)
    modoRecibo:  'cuota' | 'recibo' | 'ambos' — controls column count and headers.

      'cuota'  → 3 cols: CUOTA | FECHA DE VENCIMIENTO | IMPORTE {currency}
      'recibo' → 3 cols: NRO. DE RECIBO | FECHA DE VENCIMIENTO | IMPORTE {currency}
      'ambos'  → 4 cols: CUOTA | NRO. DE RECIBO | FECHA DE VENCIMIENTO | IMPORTE {currency}
    """
    FIN_TABLE_W_CM = 12.0
    ratio          = FIN_TABLE_W_CM / CONTENT_WIDTH_CM

    if modoRecibo == 'ambos':
        # 4-column layout — uses COL_FINANCING_AMBOS_CM scaled to FIN_TABLE_W_CM
        num_cols  = 4
        finColsCm = [round(c * ratio, 4) for c in COL_FINANCING_AMBOS_CM]
        col_headers = [
            "CUOTA",
            "NRO. DE RECIBO",
            "FECHA DE VENCIMIENTO",
            f"IMPORTE {currency}",
        ]
    else:
        # 3-column layout — uses COL_FINANCING_CM scaled to FIN_TABLE_W_CM
        num_cols  = 3
        finColsCm = [round(c * ratio, 4) for c in COL_FINANCING_CM]
        # Col 0 header changes depending on mode
        col0_header = "NRO. DE RECIBO" if modoRecibo == 'recibo' else "CUOTA"
        col_headers = [col0_header, "FECHA DE VENCIMIENTO", f"IMPORTE {currency}"]

    numTitleRows = (1 if titleRow else 0) + (1 if subtitleRow else 0)
    numRows      = numTitleRows + 1 + max(len(cuotas), 1) + 1  # +1 header, +1 total
    table        = doc.add_table(rows=numRows, cols=num_cols)
    _setTableWidth(table, width_cm=FIN_TABLE_W_CM, center=True)
    _setColumnWidths(table, finColsCm)

    rowIdx = 0

    # Optional title row (insured name or branch — merged across all cols)
    if titleRow:
        row = table.rows[rowIdx]
        row.cells[0].merge(row.cells[num_cols - 1])
        _styleCell(
            row.cells[0], titleRow, FIN_TABLE_W_CM,
            bgColor=COLOR_HEADER_BG, textColor=COLOR_HEADER_TEXT, bold=True
        )
        rowIdx += 1

    # Optional subtitle row (branch within insured — merged across all cols)
    if subtitleRow:
        row = table.rows[rowIdx]
        row.cells[0].merge(row.cells[num_cols - 1])
        _styleCell(
            row.cells[0], subtitleRow, FIN_TABLE_W_CM,
            bgColor=COLOR_HEADER_BG, textColor=COLOR_HEADER_TEXT, bold=True
        )
        rowIdx += 1

    # Column header row — dynamic per mode
    headerRow = table.rows[rowIdx]
    for col, (text, width) in enumerate(zip(col_headers, finColsCm)):
        _styleCell(
            headerRow.cells[col], text, width,
            bgColor=COLOR_HEADER_BG, textColor=COLOR_HEADER_TEXT, bold=True
        )
    _setHeaderRowInnerBorders(headerRow, num_cols)
    rowIdx += 1

    # Data rows — field mapping varies by mode
    for quota in (cuotas or [{}]):
        dataRow = table.rows[rowIdx]
        if modoRecibo == 'ambos':
            # 4-col: nro | recibo | vencimiento | importe
            _styleCell(dataRow.cells[0], str(quota.get('nro',         '')), finColsCm[0])
            _styleCell(dataRow.cells[1],     quota.get('recibo',      ''),  finColsCm[1])
            _styleCell(dataRow.cells[2],     quota.get('vencimiento', ''),  finColsCm[2])
            _styleCell(dataRow.cells[3],     quota.get('importe',     ''),  finColsCm[3])
        else:
            # 3-col: nro (or recibo) | vencimiento | importe
            _styleCell(dataRow.cells[0], str(quota.get('nro',         '')), finColsCm[0])
            _styleCell(dataRow.cells[1],     quota.get('vencimiento', ''),  finColsCm[1])
            _styleCell(dataRow.cells[2],     quota.get('importe',     ''),  finColsCm[2])
        rowIdx += 1

    # Total row — last two cells show TOTAL label and amount
    # First cell(s) are left empty with partial borders
    totalRow = table.rows[rowIdx]
    if modoRecibo == 'ambos':
        # 4-col: merge first two empty cells, then TOTAL, then amount
        _styleCell(
            totalRow.cells[0], '', finColsCm[0],
            borders={'top': True, 'bottom': False, 'left': False, 'right': False}
        )
        _styleCell(
            totalRow.cells[1], '', finColsCm[1],
            borders={'top': True, 'bottom': False, 'left': False, 'right': True}
        )
        _styleCell(totalRow.cells[2], 'TOTAL',         finColsCm[2], bold=True)
        _styleCell(totalRow.cells[3], total or '0.00', finColsCm[3], bold=True)
    else:
        # 3-col: first cell empty, then TOTAL, then amount
        _styleCell(
            totalRow.cells[0], '', finColsCm[0],
            borders={'top': True, 'bottom': False, 'left': False, 'right': True}
        )
        _styleCell(totalRow.cells[1], 'TOTAL',         finColsCm[1], bold=True)
        _styleCell(totalRow.cells[2], total or '0.00', finColsCm[2], bold=True)


def _buildFinancingSubdoc(docTemplate, formData):
    """
    Build all financing tables as a docxtpl subdocument.
    Replaces {{subdocFinancingTables}} in the template.

    Reads directly from formData['financiamiento'] — the raw UUID-keyed dict
    enriched with insuredName and policyBranch by app.py._enrichFinanciamientoNames().

    modoRecibo is read from formData and forwarded to _buildOneFinancingTable:
      'cuota'  → 3-col table, col 0 header = CUOTA
      'recibo' → 3-col table, col 0 header = NRO. DE RECIBO
      'ambos'  → 4-col table: CUOTA | NRO. DE RECIBO | FECHA DE VENCIMIENTO | IMPORTE
    """
    subdoc        = docTemplate.new_subdoc()
    currency      = formData.get('currency', '')
    tipoFin       = formData.get('tipoFinanciamiento', '')
    multiPol      = formData.get('multiplePolicies', '') == 'Varias Polizas'
    multiIns      = formData.get('multipleClients',  '') == 'Si'
    financiamento = formData.get('financiamiento', {})
    # Read the global quota mode set by the user on page 4
    modoRecibo    = formData.get('modoRecibo', 'cuota')

    if not financiamento:
        return subdoc

    # Case 1: single policy + single insured — one table, no title row
    if not multiPol and not multiIns:
        td = next(iter(financiamento.values()))
        _buildOneFinancingTable(
            subdoc,
            td.get('cuotas', []),
            td.get('total',  '0.00'),
            currency,
            modoRecibo=modoRecibo
        )

    # Case 2 Collective: multiple policies + single insured, one grouped table
    elif multiPol and not multiIns and tipoFin == 'Colectivo':
        td = next(iter(financiamento.values()))
        _buildOneFinancingTable(
            subdoc,
            td.get('cuotas', []),
            td.get('total',  '0.00'),
            currency,
            modoRecibo=modoRecibo
        )

    # Case 2 Individual: multiple policies + single insured, one table per policy
    elif multiPol and not multiIns and tipoFin == 'Individual':
        for i, td in enumerate(financiamento.values()):
            if i > 0:
                _addSpacerPara(subdoc)
            _buildOneFinancingTable(
                subdoc,
                td.get('cuotas', []),
                td.get('total',  '0.00'),
                currency,
                titleRow=td.get('policyBranch', ''),
                modoRecibo=modoRecibo
            )

    # Case 3: single policy + multiple insured, one table per insured
    elif not multiPol and multiIns:
        for i, td in enumerate(financiamento.values()):
            if i > 0:
                _addSpacerPara(subdoc)
            _buildOneFinancingTable(
                subdoc,
                td.get('cuotas', []),
                td.get('total',  '0.00'),
                currency,
                titleRow=td.get('insuredName', ''),
                modoRecibo=modoRecibo
            )

    # Case 4 Collective: multiple policies + multiple insured, one table per insured
    elif multiPol and multiIns and tipoFin == 'Colectivo':
        for i, td in enumerate(financiamento.values()):
            if i > 0:
                _addSpacerPara(subdoc)
            _buildOneFinancingTable(
                subdoc,
                td.get('cuotas', []),
                td.get('total',  '0.00'),
                currency,
                titleRow=td.get('insuredName', ''),
                modoRecibo=modoRecibo
            )

    # Case 4 Individual: multiple policies + multiple insured
    # Groups by insuredName — title = insured name, subtitle = branch name
    elif multiPol and multiIns and tipoFin == 'Individual':
        groups = {}
        for td in financiamento.values():
            name = td.get('insuredName', '')
            if name not in groups:
                groups[name] = []
            groups[name].append(td)

        first = True
        for insuredName, policies in groups.items():
            for td in policies:
                if not first:
                    _addSpacerPara(subdoc)
                first = False
                _buildOneFinancingTable(
                    subdoc,
                    td.get('cuotas', []),
                    td.get('total',  '0.00'),
                    currency,
                    titleRow=insuredName,
                    subtitleRow=td.get('policyBranch', ''),
                    modoRecibo=modoRecibo
                )

    return subdoc


# ─────────────────────────────────────────────────────────────────────────────
# ANNEX TABLE BUILDER
# ─────────────────────────────────────────────────────────────────────────────

def _buildAnnexSubdoc(docTemplate, formData):
    """
    Build the annexes table as a docxtpl subdocument.
    Replaces {{subdocAnnexTable}} in the template.

    Reads from formData['attachments'] (raw app data with keys 'anexo'/'contenido')
    and remaps to 'NUMERO'/'CONTENIDO' for display.
    Structure: header row (Anexo | Contenido) + one data row per annex.
    Anexo column is centered; Contenido column is left-aligned per spec.
    """
    subdoc = docTemplate.new_subdoc()

    # Read raw attachments and remap keys for display
    rawAnnexes = formData.get('attachments', [])
    anexos = [
        {'NUMERO': a.get('anexo', ''), 'CONTENIDO': a.get('contenido', '')}
        for a in rawAnnexes
    ]

    # Add Cesiones de Derecho row when endorsement is Detallado
    if formData.get('endorsementType') == 'Detallado':
        next_letter = chr(ord('A') + len(anexos))
        anexos.append({'NUMERO': next_letter, 'CONTENIDO': 'Cesiones de Derecho'})

    # Always render the table, even if empty (header row always present)
    numRows = 1 + max(len(anexos), 1)
    table   = subdoc.add_table(rows=numRows, cols=2)
    _setTableWidth(table)
    _setColumnWidths(table, COL_ANNEX_CM)

    # Header row
    headerRow = table.rows[0]
    _styleCell(
        headerRow.cells[0], 'Anexo', COL_ANNEX_CM[0],
        bgColor=COLOR_HEADER_BG, textColor=COLOR_HEADER_TEXT,
        bold=True, allCaps=False           # title-case per spec
    )
    _styleCell(
        headerRow.cells[1], 'Contenido', COL_ANNEX_CM[1],
        bgColor=COLOR_HEADER_BG, textColor=COLOR_HEADER_TEXT,
        bold=True, allCaps=False           # title-case per spec
    )
    _setHeaderRowInnerBorders(headerRow, 2)

    # Data rows
    if anexos:
        for i, anexo in enumerate(anexos):
            row = table.rows[i + 1]
            # Anexo column — centered
            _styleCell(
                row.cells[0],
                anexo.get('NUMERO', ''),
                COL_ANNEX_CM[0],
                alignment=WD_ALIGN_PARAGRAPH.CENTER
            )
            # Contenido column — left-aligned per spec
            _styleCell(
                row.cells[1],
                anexo.get('CONTENIDO', ''),
                COL_ANNEX_CM[1],
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                allCaps=False              # preserve original casing
            )
    else:
        # Empty placeholder row when no annexes defined
        row = table.rows[1]
        _styleCell(row.cells[0], '', COL_ANNEX_CM[0])
        _styleCell(row.cells[1], '', COL_ANNEX_CM[1])

    return subdoc




def _loadSignatureImage(docTemplate, path):
    """
    Attempt to load a signature image as a docxtpl InlineImage for insertion
    into the Word document at the {{executiveSignature}} or {{unitLeaderSignature}}
    placeholder positions.

    The image is rendered at a fixed width of 4cm inside the signature table cell.
    Height is scaled proportionally by Word based on the image's aspect ratio.

    If the file does not exist, the path is empty, or any error occurs during
    loading, the function returns an empty string '' silently. docxtpl will
    substitute '' for the template variable, leaving that cell blank. This is
    intentional — a missing signature image must never stop document generation.

    Args:
        docTemplate: DocxTemplate instance (required by InlineImage constructor).
        path (str):  File path to the signature image (e.g. 'imgs/signatures/foo.png').

    Returns:
        docxtpl.InlineImage or '' — the image object ready for template injection,
        or '' if the file is unavailable or an error occurred.
    """
    if not path:
        return ''

    try:
        import os
        # Resolve the signature path through resource_path so it works both
        # in dev (relative to project root) and in the PyInstaller bundle.
        resolved = resource_path(path)
        if not os.path.isfile(resolved):
            return ''

        # InlineImage requires the DocxTemplate instance and the file path.
        # width=Cm(4) sets the rendered width to 4cm; height scales automatically.
        from docxtpl import InlineImage
        return InlineImage(docTemplate, resolved, width=Cm(4))

    except Exception:
        # Catch any unexpected error (corrupt image, wrong format, etc.)
        # and fall back to empty string so document generation continues.
        return ''


def _xmlesc(val):
    """
    XML-escape a string value for safe Jinja2/docxtpl paragraph rendering.
    docxtpl renders template.docx as XML — bare & breaks the XML parser and
    truncates text. Table cells via python-docx are not affected (they bypass
    Jinja2), but paragraph-level variables must be escaped.
    """
    if not isinstance(val, str):
        return val
    return val.replace('&', '&amp;')


def _buildTemplateContext(formData, docTemplate):
    """
    Build the complete Jinja2 context dict from formData.

    Single source of truth for variable name mapping between the application
    and template.docx. Any rename in the template only requires a change here.
    """
    ctx = {}

    # Header
    ctx['letterNumber'] = formData.get('letterNumber', '')
    ctx['barcode']      = formData.get('barcode',      '')
    ctx['companyName']  = formData.get('companyName',  '')
    ctx['address']      = formData.get('address',      '')
    ctx['contactName']  = _xmlesc(formData.get('contactName',  ''))

    # Date in Spanish
    months = {
        1:'enero', 2:'febrero', 3:'marzo', 4:'abril', 5:'mayo', 6:'junio',
        7:'julio', 8:'agosto', 9:'septiembre', 10:'octubre',
        11:'noviembre', 12:'diciembre'
    }
    today        = date.today()
    ctx['FECHA'] = f"{today.day} de {months[today.month]} de {today.year}"

    # Letter configuration
    ctx['letterType']       = formData.get('letterType',       '')
    ctx['multiplePolicies'] = formData.get('multiplePolicies', '')
    ctx['multipleClients']  = formData.get('multipleClients',  '')
    ctx['hasPayment']       = formData.get('hasPayment',       'No')
    ctx['currency']         = formData.get('currency',         '')

    # Limit of liability
    ctx['lolAmountType']  = formData.get('lolAmountType',  '')
    ctx['lolAmountValue'] = formData.get('lolAmountValue', '')

    # Unit and executive
    ctx['unitLeader']        = formData.get('unitLeader',        '')
    ctx['executiveName']     = formData.get('executiveName',     '')
    ctx['executivePosition'] = formData.get('executivePosition', '')
    ctx['executiveEmail']    = formData.get('executiveEmail',    '')
    ctx['executiveMobile']   = formData.get('executiveMobile',   '')
    ctx['executivePhone']    = formData.get('executivePhone',    '')
    ctx['unitLeaderSignature'] = _loadSignatureImage(
        docTemplate, formData.get('unitLeaderSignature', '')
    )
    ctx['executiveSignature']  = _loadSignatureImage(
        docTemplate, formData.get('executiveSignature',  '')
    )
    # Firma table — always Lider de Unidad alongside the Executive.
    # Template uses {{unitLeader}}, {{unitLeaderSignature}}, and the
    # hardcoded text "GERENTE DE UNIDAD" written directly in build_template.js.

    # Liquidacion header — varies by company, configurable in db.json
    try:
        with open(resource_path('db.json'), 'r', encoding='utf-8') as _dbf:
            _db = _dbf.read()
        import json as _json
        _liq_headers = _json.loads(_db).get('nomenclatura_comprobantes', {})
    except Exception:
        _liq_headers = {}
    _aseguradora = formData.get('POLIZA_ASEGURADORA', '')
    ctx['liquidacionHeader'] = _liq_headers.get(_aseguradora, 'LIQUIDACIÓN')

    # Policy reference fields (used in reference line, not in tables)
    ctx['POLIZA_ASEGURADORA'] = formData.get('POLIZA_ASEGURADORA', '')
    ctx['POLIZA_INICIO']      = formData.get('POLIZA_INICIO',      '')
    ctx['POLIZA_FIN']         = formData.get('POLIZA_FIN',         '')
    ctx['POLIZA_RAMO']        = _xmlesc(formData.get('POLIZA_RAMO',        ''))

    # Endorsement
    ctx['endorsementType']    = formData.get('endorsementType',    'No')   # 'Simple' | 'Detallado' | 'No'
    ctx['ENDOSO_RAMO']         = formData.get('ENDOSO_RAMO',         '')
    ctx['ENDOSO_BENEFICIARIO'] = formData.get('ENDOSO_BENEFICIARIO', '')

    # Annexes — remap {anexo, contenido} to {NUMERO, CONTENIDO}
    ctx['ANEXOS'] = [
        {'NUMERO': a.get('anexo', ''), 'CONTENIDO': a.get('contenido', '')}
        for a in formData.get('attachments', [])
    ]

    # Financing context (firstDueDate used in template intro paragraph)
    ctx = _buildFinancingContext(ctx, formData)

    # Subdocuments — complex tables built with python-docx
    ctx['subdocPolicyTables']    = _buildPolicySubdoc(docTemplate, formData, liq_header=ctx.get('liquidacionHeader', 'LIQUIDACIÓN'))
    ctx['subdocFinancingTables'] = (
        _buildFinancingSubdoc(docTemplate, formData)
        if formData.get('hasPayment') == 'Si'
        else docTemplate.new_subdoc()
    )
    ctx['subdocAnnexTable'] = _buildAnnexSubdoc(docTemplate, formData)

    return ctx


def _buildFinancingContext(ctx, formData):
    """
    Transform the raw financiamiento dict (UUID-keyed) into named lists
    for the financing subdocument builder and for firstDueDate.
    """
    financiamento = formData.get('financiamiento', {})
    tipoFin       = formData.get('tipoFinanciamiento', '')
    multiPolicies = formData.get('multiplePolicies', '') == 'Varias Polizas'
    multiInsured  = formData.get('multipleClients',  '') == 'Si'

    ctx['CUOTAS']            = []
    ctx['financingTotal']    = '0.00'
    ctx['firstDueDate']      = ''
    ctx['FRACCIONAMIENTO']   = []
    ctx['RESPONSABLES_PAGO'] = []

    if not financiamento:
        return ctx

    # Case 1: single policy + single insured
    if not multiPolicies and not multiInsured:
        tableData             = next(iter(financiamento.values()),
                                     {'cuotas': [], 'total': '0.00'})
        ctx['CUOTAS']         = tableData.get('cuotas', [])
        ctx['financingTotal'] = tableData.get('total',  '0.00')
        ctx['firstDueDate']   = _firstDueDate(ctx['CUOTAS'])

    # Case 2 Collective: multiple policies + single insured, grouped
    elif multiPolicies and not multiInsured and tipoFin == 'Colectivo':
        tableData             = next(iter(financiamento.values()),
                                     {'cuotas': [], 'total': '0.00'})
        ctx['CUOTAS']         = tableData.get('cuotas', [])
        ctx['financingTotal'] = tableData.get('total',  '0.00')
        ctx['firstDueDate']   = _firstDueDate(ctx['CUOTAS'])

    # Case 2 Individual: multiple policies + single insured, one table per policy
    elif multiPolicies and not multiInsured and tipoFin == 'Individual':
        items = [
            {'policyBranch': td.get('policyBranch', ''),
             'cuotas':       td.get('cuotas',       []),
             'total':        td.get('total',        '0.00')}
            for td in financiamento.values()
        ]
        ctx['FRACCIONAMIENTO'] = items
        ctx['firstDueDate']    = _firstDueDate(items[0]['cuotas'] if items else [])

    # Case 3: single policy + multiple insured (always per-insured)
    elif not multiPolicies and multiInsured:
        items = [
            {'insuredName': td.get('insuredName', ''),
             'cuotas':      td.get('cuotas',      []),
             'total':       td.get('total',       '0.00')}
            for td in financiamento.values()
        ]
        ctx['RESPONSABLES_PAGO'] = items
        ctx['firstDueDate']      = _firstDueDate(items[0]['cuotas'] if items else [])

    # Case 4 Collective: multiple policies + multiple insured, grouped per insured
    elif multiPolicies and multiInsured and tipoFin == 'Colectivo':
        items = [
            {'insuredName': td.get('insuredName', ''),
             'cuotas':      td.get('cuotas',      []),
             'total':       td.get('total',       '0.00')}
            for td in financiamento.values()
        ]
        ctx['RESPONSABLES_PAGO'] = items
        ctx['firstDueDate']      = _firstDueDate(items[0]['cuotas'] if items else [])

    # Case 4 Individual: multiple policies + multiple insured, one table per policy
    elif multiPolicies and multiInsured and tipoFin == 'Individual':
        groups = {}
        for key, td in financiamento.items():
            insuredId = key.split('::')[0] if '::' in key else key
            if insuredId not in groups:
                groups[insuredId] = {
                    'insuredName': td.get('insuredName', ''),
                    'polizas':     []
                }
            groups[insuredId]['polizas'].append({
                'policyBranch': td.get('policyBranch', ''),
                'cuotas':       td.get('cuotas',       []),
                'total':        td.get('total',        '0.00'),
            })
        ctx['RESPONSABLES_PAGO'] = list(groups.values())
        firstGroup = next(iter(groups.values()), {})
        ctx['firstDueDate'] = _firstDueDate(
            firstGroup.get('polizas', [{}])[0].get('cuotas', [])
            if firstGroup.get('polizas') else []
        )

    return ctx


def _firstDueDate(quotaList):
    """Return vencimiento of first quota, or '' if empty."""
    return quotaList[0].get('vencimiento', '') if quotaList else ''


# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENT GENERATION
# ─────────────────────────────────────────────────────────────────────────────

def _convertToPdf(docx_path):
    """
    Convert a .docx file to PDF using docx2pdf (requires Microsoft Word on Windows).
    The PDF is saved alongside the .docx with the same name but .pdf extension.

    Returns:
        str: path to generated PDF, or raises Exception on failure.
    """
    import os
    from docx2pdf import convert
    pdf_path = os.path.splitext(docx_path)[0] + '.pdf'
    convert(docx_path, pdf_path)
    return pdf_path


def generateDocument(formData, parentWidget):
    """
    Generate the Word dispatch letter, and optionally convert to PDF.

    outputFormat in formData controls output:
      'word' — save .docx only
      'pdf'  — save .pdf only  (docx generated internally then deleted)
      'both' — save both .docx and .pdf

    Returns:
        dict: {'success': bool, 'path': str, 'pdf_path': str, 'error': str}
    """
    import os
    try:
        outputFormat = formData.get('outputFormat', 'word')

        # Choose dialog filter and default filename based on format
        if outputFormat == 'pdf':
            dialogFilter  = "PDF (*.pdf)"
            defaultName   = "carta_despacho.pdf"
        else:
            dialogFilter  = "Word (*.docx)"
            defaultName   = "carta_despacho.docx"

        savePath, _ = QFileDialog.getSaveFileName(
            parentWidget,
            "Guardar Documento",
            defaultName,
            dialogFilter,
        )

        if not savePath:
            return {'success': False, 'error': 'Cancelado por usuario', 'path': '', 'pdf_path': ''}

        # For PDF-only mode, generate docx to a temp path then convert
        if outputFormat == 'pdf':
            # Ensure the save path ends in .pdf
            if not savePath.lower().endswith('.pdf'):
                savePath += '.pdf'
            temp_docx = os.path.splitext(savePath)[0] + '_temp.docx'
        else:
            # Ensure the save path ends in .docx
            if not savePath.lower().endswith('.docx'):
                savePath += '.docx'
            temp_docx = None

        docx_path = temp_docx if outputFormat == 'pdf' else savePath

        # Load template before building context — subdocuments need the reference
        docTemplate = DocxTemplate(resource_path('template.docx'))
        context     = _buildTemplateContext(formData, docTemplate)
        docTemplate.render(context)
        docTemplate.save(docx_path)

        pdf_path = ''

        if outputFormat == 'pdf':
            pdf_path = _convertToPdf(docx_path)
            # Rename converted pdf to the user-chosen path
            if pdf_path != savePath:
                if os.path.exists(savePath):
                    os.remove(savePath)
                os.rename(pdf_path, savePath)
                pdf_path = savePath
            # Remove temp docx
            if os.path.exists(docx_path):
                os.remove(docx_path)
            return {'success': True, 'path': savePath, 'pdf_path': pdf_path, 'error': ''}

        elif outputFormat == 'both':
            pdf_path = _convertToPdf(savePath)

        return {'success': True, 'path': savePath, 'pdf_path': pdf_path, 'error': ''}

    except FileNotFoundError:
        return {'success': False, 'path': '', 'pdf_path': '',
                'error': 'Archivo template.docx no encontrado'}
    except Exception as error:
        return {'success': False, 'path': '', 'pdf_path': '', 'error': str(error)}

def generateManualSiniestros(formData, selectedItems, savePath):
    """
    Generate the Manual de Procedimientos en Caso de Siniestros Word document.

    Loads template_siniestros.docx, builds the Jinja2 context from formData
    and selectedItems, renders and saves to savePath.

    Args:
        formData:      Complete dict from all pages in app.py
        selectedItems: List of dicts [{number, name, startPage}] from PageManual
        savePath:      Full path where the .docx will be saved

    Returns:
        dict: {'success': bool, 'path': str, 'error': str}
    """
    try:
        docTemplate = DocxTemplate(resource_path('template_siniestros.docx'))

        # Build show_item flags for items 2-13
        selected_numbers = {item['number'] for item in selectedItems}
        show_flags = {f'show_item_{n}': (n in selected_numbers) for n in range(2, 14)}

        # Total body pages: each item is 2 pages except item 13 which is 3
        total_pages = sum(3 if item['number'] == 13 else 2 for item in selectedItems)
        # Item 1 (Aspectos Generales) always contributes 2 pages
        total_pages += 2

        context = {
            # Company / document info
            'companyName':     formData.get('companyName',     ''),
            # Unit leader
            'unitLeader':      formData.get('unitLeader',      ''),
            'leaderMobile':    formData.get('leaderMobile',    ''),
            'leaderEmail':     formData.get('leaderEmail',     ''),
            # Executive
            'executiveName':   formData.get('executiveName',   ''),
            'executiveMobile': formData.get('executiveMobile', ''),
            'executiveEmail':  formData.get('executiveEmail',  ''),
            # TOC
            'indexItems':      selectedItems,
            'total_pages':     total_pages,
            # Section visibility flags
            **show_flags,
        }

        docTemplate.render(context)
        docTemplate.save(savePath)

        return {'success': True, 'path': savePath, 'error': ''}

    except FileNotFoundError:
        return {'success': False, 'path': '',
                'error': 'Archivo template_siniestros.docx no encontrado'}
    except Exception as error:
        return {'success': False, 'path': '', 'error': str(error)}



# ─────────────────────────────────────────────────────────────────────────────
# GARANTÍAS PARTICULARES — direct python-docx builder
#
# Uses GARANTIAS_PARTICULARES.docx (the original uploaded template) as the
# base document. All styles and numbering definitions are already present in
# that file — no JS-generated template needed for the content section.
#
# Verified numIds in GARANTIAS_PARTICULARES.docx:
#   numId=1  → decimal auto-counter (SectionStart style uses this internally)
#   numId=2  → » bullet, Noto Sans, left=227 hanging=227
#   numId=3  → decimal %1. list, Noto Sans, left=227 hanging=227 (same indent as chevron)
#
# Verified styleIds:
#   SectionStart → auto-number + pageBreakBefore + blue border (all in style def)
#   Heading      → Noto Sans bold 18pt
#   TextBody     → Noto Sans 11pt, justified, 0 spacing
# ─────────────────────────────────────────────────────────────────────────────

def _parseGuaranteesHtml(html):
    """
    Parse QTextEdit HTML into a list of paragraph dicts for Word generation.

    Returns: [{'type': 'bullet'|'plain', 'runs': [{'text': str, 'bold': bool,
               'italic': bool, 'underline': bool}, ...]}, ...]

    QTextEdit generates spans with inline styles, not semantic <b>/<i>/<u> tags:
      font-weight: 600/700/bold   → bold
      font-style: italic           → italic
      text-decoration: underline   → underline

    A span stack tracks exactly which attributes each span opened so they
    are correctly undone on </span> without cross-contamination between
    paragraphs or sibling spans.

    Blank paragraphs (<p></p>) are preserved as NBSP.
    """
    from html.parser import HTMLParser

    class _Parser(HTMLParser):
        def __init__(self):
            super().__init__()
            self.paragraphs  = []
            self._in_bullet  = False
            self._para       = None
            # Current active formatting (toggled by tags/spans)
            self._bold       = False
            self._italic     = False
            self._underline  = False
            # Stack: each entry is (bold_delta, italic_delta, underline_delta)
            # so we know exactly what to undo when the matching </span> arrives
            self._span_stack = []

        def _push_text(self, text):
            if self._para is None or not text:
                return
            runs = self._para['runs']
            fmt  = (self._bold, self._italic, self._underline)
            if runs and (runs[-1]['bold'], runs[-1]['italic'], runs[-1]['underline']) == fmt:
                runs[-1]['text'] += text
            else:
                runs.append({'text': text, 'bold': self._bold,
                             'italic': self._italic, 'underline': self._underline})

        def handle_starttag(self, tag, attrs):
            attrs_dict = dict(attrs)
            if tag == 'ul':
                self._in_bullet = True
            elif tag in ('li', 'p'):
                if self._in_bullet:
                    para_type = 'bullet'
                else:
                    # Detect residual list indent: QTextEdit adds margin-left
                    # to plain paragraphs that were de-bulleted while retaining indent
                    style = attrs_dict.get('style', '')
                    import re as _re2
                    has_indent = bool(_re2.search('margin-left' + r'\s*' + ':' + r'\s*' + '[1-9]', style))
                    para_type = 'indented' if has_indent else 'plain'
                self._para = {'type': para_type, 'runs': []}
            elif tag in ('b', 'strong'):
                self._bold = True
            elif tag in ('i', 'em'):
                self._italic = True
            elif tag == 'u':
                self._underline = True
            elif tag == 'span':
                style = attrs_dict.get('style', '')
                # Track exactly what this span activates so </span> can undo it
                opened_bold = opened_italic = opened_underline = False

                if 'font-weight' in style:
                    # QTextEdit uses numeric weights (600, 700) or 'bold'
                    import re as _re
                    weight_match = _re.search(r'font-weight\s*:\s*(\w+)', style)
                    if weight_match:
                        val = weight_match.group(1)
                        if val == 'bold' or (val.isdigit() and int(val) >= 600):
                            if not self._bold:
                                self._bold = True
                                opened_bold = True

                if 'font-style' in style and 'italic' in style:
                    if not self._italic:
                        self._italic = True
                        opened_italic = True

                if 'text-decoration' in style and 'underline' in style:
                    if not self._underline:
                        self._underline = True
                        opened_underline = True

                self._span_stack.append((opened_bold, opened_italic, opened_underline))

        def handle_endtag(self, tag):
            if tag == 'ul':
                self._in_bullet = False
            elif tag in ('li', 'p'):
                if self._para is not None:
                    runs      = self._para['runs']
                    total_txt = ''.join(r['text'] for r in runs).strip()
                    if total_txt:
                        self.paragraphs.append(self._para)
                    elif self._para['type'] == 'plain':
                        self.paragraphs.append({
                            'type': 'plain',
                            'runs': [{'text': '\u00A0', 'bold': False,
                                      'italic': False, 'underline': False}],
                        })
                self._para = None
                # Reset formatting at paragraph boundary — QTextEdit scopes
                # all formatting within each paragraph's own span tree
                self._bold = self._italic = self._underline = False
                self._span_stack.clear()
            elif tag in ('b', 'strong'):
                self._bold = False
            elif tag in ('i', 'em'):
                self._italic = False
            elif tag == 'u':
                self._underline = False
            elif tag == 'span':
                if self._span_stack:
                    ob, oi, ou = self._span_stack.pop()
                    if ob:
                        self._bold = False
                    if oi:
                        self._italic = False
                    if ou:
                        self._underline = False

        def handle_data(self, data):
            self._push_text(data)

    parser = _Parser()
    parser.feed(html or '')
    return parser.paragraphs


def generateGuarantees(formData, guaranteesData, savePath):
    """
    Generate the Garantías Particulares Word document.

    Opens GARANTIAS_PARTICULARES.docx directly with python-docx (no docxtpl,
    no subdocuments). Replaces {{companyName}} in the cover, clears content
    paragraphs after the section break, and writes policy sections directly.

    All styles (SectionStart, Heading, TextBody) and numbering (numId 2 for
    chevron bullets, numId 5 for decimal numbered lists) are resolved from the
    template file where they are defined.

    Args:
        formData:       Complete dict from all pages in app.py.
        guaranteesData: {widgetId: {'branch': str, 'numero': str, 'html': str}}
        savePath:       Full path where the .docx will be saved.

    Returns:
        dict: {'success': bool, 'path': str, 'error': str}
    """
    try:
        import os
        from docx import Document as _DocxDoc
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _Oxm

        # ── locate template ───────────────────────────────────────────────────
        # Uses template_garantias.docx — generated by build_template_garantias.js.
        # This file carries all custom styles (SectionStart, Heading, TextBody)
        # and numbering definitions (numId 1/2/5) needed for content generation.
        from strings import S as _S
        template_path = resource_path('template_garantias.docx')
        if not os.path.exists(template_path):
            raise FileNotFoundError(_S['err_template_garantias'])

        doc  = _DocxDoc(template_path)
        body = doc.element.body

        # ── helpers ───────────────────────────────────────────────────────────
        def _set_style(para, style_id):
            """Inject w:pStyle by ID — bypasses python-docx registry lookup."""
            pPr    = para._p.get_or_add_pPr()
            pStyle = pPr.find(_qn('w:pStyle'))
            if pStyle is None:
                pStyle = _Oxm('w:pStyle')
                pPr.insert(0, pStyle)
            pStyle.set(_qn('w:val'), style_id)

        def _set_numbering(para, num_id_val):
            """Inject w:numPr for list formatting."""
            pPr   = para._p.get_or_add_pPr()
            existing = pPr.find(_qn('w:numPr'))
            if existing is not None:
                pPr.remove(existing)
            numPr = _Oxm('w:numPr')
            ilvl  = _Oxm('w:ilvl'); ilvl.set(_qn('w:val'), '0')
            numId = _Oxm('w:numId'); numId.set(_qn('w:val'), str(num_id_val))
            numPr.append(ilvl); numPr.append(numId)
            pPr.append(numPr)

        def _set_indent(para, left_dxa, hanging_dxa):
            """Set paragraph indent via XML."""
            pPr = para._p.get_or_add_pPr()
            existing = pPr.find(_qn('w:ind'))
            if existing is not None:
                pPr.remove(existing)
            ind = _Oxm('w:ind')
            ind.set(_qn('w:left'),    str(left_dxa))
            ind.set(_qn('w:hanging'), str(hanging_dxa))
            pPr.append(ind)

        def _zero_spacing(para):
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after  = Pt(0)

        def _noto_run(para, text, size_pt=11, bold=False):
            """Add a run with explicit Noto Sans font and size."""
            run = para.add_run(text)
            run.font.name = 'Noto Sans'
            run.font.size = Pt(size_pt)
            run.font.bold = bold
            return run

        # Indent values for bullet paragraphs.
        # Word dialog shows: Left = w:left - w:hanging, By = w:hanging
        # To display Left=0.6cm, By=0.4cm → w:left = 0.6+0.4 = 1.0cm = 567 DXA
        LEFT_DXA    = 567   # 1.0cm → Word displays as Left: 0.6cm
        HANGING_DXA = 227   # 0.4cm → Word displays as By: 0.4cm

        # ── Step 1: replace {{companyName}} in cover ──────────────────────────
        # The placeholder may be split across multiple runs in the XML (Word
        # splits inline text at arbitrary boundaries). Replace by rebuilding the
        # full paragraph text, then writing it into the first run and clearing
        # the rest — only for paragraphs that contain the placeholder.
        company_name = formData.get('companyName', '')
        for para in doc.paragraphs:
            full_text = para.text
            if '{{companyName}}' in full_text or '{{ companyName }}' in full_text:
                new_text = (full_text
                            .replace('{{companyName}}',   company_name)
                            .replace('{{ companyName }}', company_name))
                # Preserve the formatting of the first run; clear subsequent runs
                runs = para.runs
                if runs:
                    runs[0].text = new_text
                    for run in runs[1:]:
                        run.text = ''

        # ── Step 2: find section break + clear content paragraphs ─────────────
        # The section break (w:sectPr inside w:pPr) separates cover from content.
        # All paragraphs after it (content section) get removed and rebuilt.
        all_paras_el = body.findall(_qn('w:p'))
        section_break_idx = None
        for i, p_el in enumerate(all_paras_el):
            pPr = p_el.find(_qn('w:pPr'))
            if pPr is not None and pPr.find(_qn('w:sectPr')) is not None:
                section_break_idx = i
                break

        if section_break_idx is None:
            raise ValueError('Section break not found in GARANTIAS_PARTICULARES.docx')

        # Remove content paragraphs (after section break), keep final body sectPr
        content_paras = all_paras_el[section_break_idx + 1:]
        body_sectPr   = body.find(_qn('w:sectPr'))  # final document sectPr

        for p_el in content_paras:
            body.remove(p_el)

        # ── Step 3: insert new paragraphs before final sectPr ─────────────────
        def _new_para():
            """Create and insert a new paragraph element before body sectPr."""
            from docx.text.paragraph import Paragraph as _Para
            p_el = _Oxm('w:p')
            body.insert(list(body).index(body_sectPr), p_el)
            return _Para(p_el, doc)

        # ── Step 4: build content ─────────────────────────────────────────────
        entries = list(guaranteesData.items())

        for wid, entry in entries:
            branch = entry.get('branch', '')
            numero = entry.get('numero', '')
            html   = entry.get('html',   '')
            paras  = _parseGuaranteesHtml(html)

            # Section Start paragraph
            # styleId='SectionStart' carries numId=1 (decimal auto-counter),
            # pageBreakBefore, and blue bottom border — all in the style definition.
            # No extra w:numPr needed. NBSP run prevents Times New Roman fallback.
            sec = _new_para()
            _set_style(sec, 'SectionStart')
            _noto_run(sec, '\u00A0', size_pt=36)

            # Heading: "PÓLIZA {BRANCH} {NUMERO}"
            hdg = _new_para()
            _set_style(hdg, 'Heading')
            _noto_run(hdg, f"PÓLIZA {branch.upper()} {numero}".strip(), size_pt=18, bold=True)
            _zero_spacing(hdg)

            # Spacer — NBSP Noto Sans 11pt between heading and content
            sp = _new_para()
            _noto_run(sp, '\u00A0', size_pt=11)
            _zero_spacing(sp)

            # Content paragraphs from QTextEdit HTML
            # Each item has 'type' and 'runs' (list of {text, bold, italic, underline})
            for item in paras:
                item_type = item['type']
                runs      = item['runs']
                is_blank  = (len(runs) == 1 and runs[0]['text'] == '\u00A0')

                c = _new_para()
                _zero_spacing(c)

                if item_type == 'bullet':
                    _set_style(c, 'BulletBody')
                    _set_numbering(c, 2)
                    _set_indent(c, LEFT_DXA, HANGING_DXA)
                    c.paragraph_format.space_after = Pt(6)

                elif item_type == 'indented':
                    # Paragraph that had its bullet removed but kept the visual indent.
                    # Maps to IndentedBody style so the docx preserves the indent.
                    _set_style(c, 'IndentedBody')
                    _set_indent(c, LEFT_DXA, 0)
                    c.paragraph_format.space_after = Pt(6)

                elif is_blank:
                    # Blank paragraph — no spacing, contextualSpacing active
                    _set_style(c, 'TextBody')
                    pPr = c._p.get_or_add_pPr()
                    ctx = _Oxm('w:contextualSpacing')
                    ctx.set(_qn('w:val'), '1')
                    existing = pPr.find(_qn('w:contextualSpacing'))
                    if existing is not None:
                        pPr.remove(existing)
                    pPr.append(ctx)

                else:
                    # Plain text — TextBody, after=6pt
                    _set_style(c, 'TextBody')
                    c.paragraph_format.space_after = Pt(6)

                # Add runs with inline formatting (bold/italic/underline)
                for run_data in runs:
                    run            = c.add_run(run_data['text'])
                    run.font.name  = 'Noto Sans'
                    run.font.size  = Pt(11)
                    run.bold       = run_data['bold']      or None
                    run.italic     = run_data['italic']    or None
                    run.underline  = run_data['underline'] or None

        # ── Step 5: save ──────────────────────────────────────────────────────
        doc.save(savePath)
        return {'success': True, 'path': savePath, 'error': ''}

    except FileNotFoundError as fe:
        return {'success': False, 'path': '', 'error': str(fe)}
    except Exception as error:
        return {'success': False, 'path': '', 'error': str(error)}


# ─────────────────────────────────────────────────────────────────────────────
# CESIONES DE DERECHO — direct python-docx builder
# Uses template_cesiones.docx (generated by build_template_cesiones.js).
# ─────────────────────────────────────────────────────────────────────────────

def generateCesiones(formData, endorsementData, savePath, insuredNames=None, policyNumbers=None):
    """
    Generate the Cesiones de Derecho Word document.

    endorsementData: { policyKey: {'checked': bool, 'rows': [...] | 'groups': [...]} }
    Each key maps to either a CesionesWidget result (rows) or TrecWidget result (groups).

    insuredNames: { widgetId: insuredName } - mapping for insured headers in Word
    policyNumbers: { policyKey: policyNumber } - fresh policy numbers from pagePolicies

    Returns:
        dict: {'success': bool, 'path': str, 'error': str}
    """
    try:
        import os
        from docx import Document as _DocxDoc
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _Oxm

        template_path = resource_path('template_cesiones.docx')
        if not os.path.exists(template_path):
            return {'success': False, 'path': '', 'error': 'template_cesiones.docx no encontrado'}

        doc  = _DocxDoc(template_path)
        body = doc.element.body

        # ── Replace {{companyName}} in cover ──────────────────────────────────
        company_name = formData.get('companyName', '')
        for para in doc.paragraphs:
            full_text = para.text
            if '{{companyName}}' in full_text or '{{ companyName }}' in full_text:
                new_text = (full_text
                            .replace('{{companyName}}',   company_name)
                            .replace('{{ companyName }}', company_name))
                runs = para.runs
                if runs:
                    runs[0].text = new_text
                    for run in runs[1:]:
                        run.text = ''

        # ── Find section break and clear content ──────────────────────────────
        all_paras_el = body.findall(_qn('w:p'))
        section_break_idx = None
        for i, p_el in enumerate(all_paras_el):
            pPr = p_el.find(_qn('w:pPr'))
            if pPr is not None and pPr.find(_qn('w:sectPr')) is not None:
                section_break_idx = i
                break

        if section_break_idx is None:
            return {'success': False, 'path': '', 'error': 'Section break not found in template_cesiones.docx'}

        content_paras = all_paras_el[section_break_idx + 1:]
        body_sectPr   = body.find(_qn('w:sectPr'))
        for p_el in content_paras:
            body.remove(p_el)

        # ── Helpers ───────────────────────────────────────────────────────────
        def _new_para():
            from docx.text.paragraph import Paragraph as _Para
            p_el = _Oxm('w:p')
            body.insert(list(body).index(body_sectPr), p_el)
            return _Para(p_el, doc)

        def _set_style(para, style_id):
            pPr    = para._p.get_or_add_pPr()
            pStyle = pPr.find(_qn('w:pStyle'))
            if pStyle is None:
                pStyle = _Oxm('w:pStyle')
                pPr.insert(0, pStyle)
            pStyle.set(_qn('w:val'), style_id)

        def _add_page_break():
            """Insert explicit page break after table."""
            p_el = _Oxm('w:p')
            r_el = _Oxm('w:r')
            br_el = _Oxm('w:br')
            br_el.set(_qn('w:type'), 'page')
            r_el.append(br_el)
            p_el.append(r_el)
            body.insert(list(body).index(body_sectPr), p_el)

        def _add_spacer(element_after, font_val):
            """Insert ZeroWidthSpace paragraph AFTER element (e.g., table).
            Spacing After 0, Line Spacing Single."""
            sep_el = _Oxm('w:p')
            r_el = _Oxm('w:r')
            t_el = _Oxm('w:t')
            t_el.text = '\u200B'
            rPr_el = _Oxm('w:rPr')
            sz_el = _Oxm('w:sz')
            sz_el.set(_qn('w:val'), str(font_val))
            rPr_el.append(sz_el)
            rFonts_el = _Oxm('w:rFonts')
            rFonts_el.set(_qn('w:ascii'), 'Noto Sans')
            rFonts_el.set(_qn('w:hAnsi'), 'Noto Sans')
            rPr_el.append(rFonts_el)
            r_el.append(rPr_el)
            r_el.append(t_el)
            sep_el.append(r_el)
            pPr_el = _Oxm('w:pPr')
            spacing_el = _Oxm('w:spacing')
            spacing_el.set(_qn('w:after'), '0')
            pPr_el.append(spacing_el)
            sep_el.insert(0, pPr_el)
            element_after.addnext(sep_el)

        def _add_para_after(element_after, font_size_pt):
            """Insert paragraph with specified font size AFTER a given element (e.g., table)."""
            from docx.text.paragraph import Paragraph as _Para
            p_el = _Oxm('w:p')
            r_el = _Oxm('w:r')
            t_el = _Oxm('w:t')
            t_el.text = '\u00A0'
            r_el.append(t_el)
            p_el.append(r_el)
            element_after.addnext(p_el)
            para = _Para(p_el, doc)
            run = para.runs[0]
            run.font.size = Pt(font_size_pt)
            run.font.name = 'Noto Sans'
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1
            return para

        def _zero_spacing(para):
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after  = Pt(0)

        def _set_spacing(para, before_pt=0, after_pt=0):
            """Set paragraph spacing."""
            if before_pt is not None:
                para.paragraph_format.space_before = Pt(before_pt)
            if after_pt is not None:
                para.paragraph_format.space_after = Pt(after_pt)

        def _noto_run(para, text, size_pt=11, bold=False, color=None):
            run = para.add_run(text)
            run.font.name = 'Noto Sans'
            run.font.size = Pt(size_pt)
            run.font.bold = bold
            if color:
                run.font.color.rgb = color
            return run

        def _set_cell_shading(cell, fill_color, border_color='000000', add_borders=True, edges=None):
            """Set cell background color (shading). fill_color: '000F47' for header, 'FFFFFF' for data.
            add_borders=False for data rows (no borders).
            edges: list of ['top','bottom','left','right'] to apply specific borders, or None for all."""
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shading = _Oxm('w:shd')
            shading.set(_qn('w:fill'), fill_color)
            shading.set(_qn('w:val'), 'clear')
            tcPr.append(shading)
            
            if add_borders:
                # Add borders only to header cells
                tcBorders = _Oxm('w:tcBorders')
                edges_to_apply = edges if edges else ['top', 'bottom', 'left', 'right']
                for edge in edges_to_apply:
                    edge_el = _Oxm(f'w:{edge}')
                    edge_el.set(_qn('w:val'), 'single')
                    edge_el.set(_qn('w:sz'), '4')
                    edge_el.set(_qn('w:color'), border_color)
                    tcBorders.append(edge_el)
                tcPr.append(tcBorders)

        def _set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
            """Set specific cell borders by side. Colors: '000F47' for navy, 'FFFFFF' for white."""
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = _Oxm('w:tcBorders')
            edges = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
            for edge_name, color in edges.items():
                if color:
                    edge_el = _Oxm(f'w:{edge_name}')
                    edge_el.set(_qn('w:val'), 'single')
                    edge_el.set(_qn('w:sz'), '4')
                    edge_el.set(_qn('w:color'), color)
                    tcBorders.append(edge_el)
            tcPr.append(tcBorders)

        def _set_para_shading(para, fill_color):
            """Set paragraph background color (shading)."""
            pPr = para._element.get_or_add_pPr()
            shading = _Oxm('w:shd')
            shading.set(_qn('w:fill'), fill_color)
            shading.set(_qn('w:val'), 'clear')
            pPr.append(shading)

        def _set_cell_text_color(cell, color_hex):
            """Set cell text color."""
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = color_hex

        def _add_table_dynamic(branch, numero, columns, rows, currency, endosatario='', hasMultipleInsured=False):
            """Add a dynamic-column table with styled header. Returns (True, table_element) or False."""
            if not rows:
                return False

            num_cols = len(columns) if columns else 3

            if endosatario:
                tbl = doc.add_table(rows=2 + len(rows), cols=num_cols)
            else:
                tbl = doc.add_table(rows=1 + len(rows), cols=num_cols)

            tbl._element.getparent().remove(tbl._element)
            body.insert(list(body).index(body_sectPr), tbl._element)
            table_element = tbl._element

            current_row_idx = 0

            if endosatario:
                erow = tbl.rows[0]
                erun = erow.cells[0].paragraphs[0].runs[0] if erow.cells[0].paragraphs[0].runs else erow.cells[0].paragraphs[0].add_run()
                erun.text = endosatario
                erun.font.name = 'Noto Sans'
                erun.font.size = Pt(11)
                erun.font.bold = True
                erun.font.color.rgb = RGBColor(0xCD, 0xEC, 0xFF)
                erow.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                _set_spacing(erow.cells[0].paragraphs[0], 3, 3)
                for ci in range(1, num_cols):
                    erow.cells[ci].merge(erow.cells[0])
                erow.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                _set_cell_shading(erow.cells[0], '000F47', 'FFFFFF')
                _set_cell_borders(erow.cells[0], top='000F47', left='000F47', right='000F47', bottom='FFFFFF')
                current_row_idx = 1

            hrow = tbl.rows[current_row_idx]
            for i, h in enumerate(columns):
                cell = hrow.cells[i]
                cell.text = h
                run = cell.paragraphs[0].runs[0]
                run.font.name = 'Noto Sans'
                run.font.size = Pt(10)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xCD, 0xEC, 0xFF)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                _set_spacing(cell.paragraphs[0], 3, 3)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                _set_cell_shading(cell, '000F47', 'FFFFFF')
                if i == 0:
                    _set_cell_borders(cell, left='000F47', right='FFFFFF', bottom='000F47')
                elif i == len(columns) - 1:
                    _set_cell_borders(cell, left='FFFFFF', right='000F47', bottom='000F47')
                else:
                    _set_cell_borders(cell, left='FFFFFF', right='FFFFFF', bottom='000F47')

            for ri, row in enumerate(rows):
                dr = tbl.rows[ri + 1 + current_row_idx]
                _set_cell_shading(dr.cells[0], 'FFFFFF', '000F47', True)
                if isinstance(row, dict):
                    for ci in range(num_cols):
                        col_name = columns[ci].lower().replace(' ', '_')
                        value = row.get(col_name, row.get(ci, ''))
                        dr.cells[ci].text = str(value) if value else ''
                elif isinstance(row, (list, tuple)):
                    for ci in range(num_cols):
                        value = row[ci] if ci < len(row) else ''
                        dr.cells[ci].text = str(value) if value else ''
                for ci in range(num_cols):
                    _set_cell_shading(dr.cells[ci], 'FFFFFF', '000F47', True)
                    dr.cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _set_spacing(dr.cells[ci].paragraphs[0], 3, 3)
                    dr.cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    for run in dr.cells[ci].paragraphs[0].runs:
                        run.font.name = 'Noto Sans'
                        run.font.size = Pt(10)

            return True, table_element

        def _add_table_3col(branch, numero, rows, currency, hasMultipleInsured=False):
            columns = ['Detalle', f'Monto Endosado {currency}', 'Endosatario']
            return _add_table_dynamic(branch, numero, columns, rows, currency, hasMultipleInsured=hasMultipleInsured)

        def _add_table_by_endorsee(branch, numero, groups, currency, hasMultipleInsured=False):
            if not groups:
                return False

            any_tables = False
            if len(groups) == 1 and not groups[0].get('name', ''):
                cols = groups[0].get('columns', [])
                rows = groups[0].get('rows', [])
                endosatario = groups[0].get('name', '')
                any_tables = _add_table_dynamic(branch, numero, cols, rows, currency, endosatario=endosatario, hasMultipleInsured=hasMultipleInsured)
                return any_tables

            for idx, grp in enumerate(groups):
                grp_name = grp.get('name', '')
                cols = grp.get('columns', [])
                rows = grp.get('rows', [])
                if not rows:
                    continue
                result = _add_table_dynamic(branch, numero, cols, rows, currency, endosatario=grp_name, hasMultipleInsured=hasMultipleInsured)
                tbl_added = result[0] if isinstance(result, tuple) else result
                if tbl_added:
                    any_tables = True
                    tables_in_body = body.findall(_qn('w:tbl'))
                    if tables_in_body and idx < len(groups) - 1:
                        _add_spacer(tables_in_body[-1], 16)

            return any_tables

        def _add_table_trec(branch, numero, groups, currency, hasMultipleInsured=False):
            """Add per-group 8-column TREC tables with styled header. Returns True if tables added."""
            # Blank paragraph: 3pt font
            sp = _new_para()
            _noto_run(sp, '\u00A0', size_pt=3)
            _zero_spacing(sp)

            any_tables = False
            for grp in groups:
                grp_name = grp.get('name', '')
                rows     = grp.get('rows', [])
                if not rows:
                    continue
                any_tables = True

                # Group sub-heading
                gh = _new_para()
                _noto_run(gh, grp_name, size_pt=12, bold=True)
                _zero_spacing(gh)

                cols = ['Equipo','Marca','Modelo','Serie','Placa','Año','Nro. Leasing',f'Suma Asegurada {currency}']
                keys = ['equipo','marca','modelo','serie','placa','anio','leasing','suma']

                # TREC table without style - borders via shading
                tbl = doc.add_table(rows=1 + len(rows), cols=8)
                tbl._element.getparent().remove(tbl._element)
                body.insert(list(body).index(body_sectPr), tbl._element)

                hrow = tbl.rows[0]
                for i, h in enumerate(cols):
                    c = hrow.cells[i]
                    c.text = h
                    run = c.paragraphs[0].runs[0]
                    run.font.name = 'Noto Sans'
                    run.font.size = Pt(9)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xCD, 0xEC, 0xFF)
                    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _set_spacing(c.paragraphs[0], 3, 3)
                    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    _set_cell_shading(c, '000F47', 'FFFFFF')

                # Data rows - white background with navy borders
                for ri, row in enumerate(rows):
                    dr = tbl.rows[ri + 1]
                    for i, k in enumerate(keys):
                        dr.cells[i].text = str(row.get(k,''))
                        _set_cell_shading(dr.cells[i], 'FFFFFF', '000F47', True)
                        dr.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        _set_spacing(dr.cells[i].paragraphs[0], 3, 3)
                        dr.cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        for run in dr.cells[i].paragraphs[0].runs:
                            run.font.name = 'Noto Sans'
                            run.font.size = Pt(9)

        # ── Build content ─────────────────────────────────────────────────────
        currency = formData.get('currency', 'S/.')
        if insuredNames is None:
            insuredNames = {}
        if policyNumbers is None:
            policyNumbers = {}

        # Group endorsement data by insured (using widgetId from key)
        # Key format: 'insuredId::policyId'
        insured_data = {}  # {insuredId: [{key, entry}, ...]}

        for key, entry in endorsementData.items():
            if not isinstance(entry, dict):
                continue
            is_checked = entry.get('checked', True)
            if is_checked is False:
                continue
            branch = entry.get('branch', '')
            if not branch:
                continue
            has_rows = isinstance(entry.get('rows'), list) and len(entry.get('rows', [])) > 0
            has_groups = isinstance(entry.get('groups'), list) and len(entry.get('groups', [])) > 0
            if not (has_rows or has_groups):
                continue

            # Extract insuredId from key (format: 'insuredId::policyId')
            insured_id = key.split('::', 1)[0] if '::' in key else key
            if insured_id not in insured_data:
                insured_data[insured_id] = []
            insured_data[insured_id].append({'key': key, 'entry': entry})

# Build content grouped by insured
        for insured_id, insured_entries in insured_data.items():
            insured_name = insuredNames.get(insured_id, '')
            has_multi_insured = len(insured_data) > 1

            # Detect mode from first entry: 'by_endorsee' = Individual, else Agrupado
            first_entry = insured_entries[0]['entry'] if insured_entries else {}
            mode = first_entry.get('mode', 'single')
            is_individual = mode == 'by_endorsee'

            # Group entries by policy_id for proper block rendering
            policies_map = {}
            for item in insured_entries:
                entry = item['entry']
                key = item['key']
                policy_id = key.split('::', 1)[1] if '::' in key else ''
                if policy_id not in policies_map:
                    policies_map[policy_id] = []
                policies_map[policy_id].append({'key': key, 'entry': entry})

            # SectionStart for ALL cases (spacing after: 12pt, line spacing single)
            sec = _new_para()
            _set_style(sec, 'SectionStart')
            _noto_run(sec, '\u00A0', size_pt=36)
            _set_spacing(sec, None, 12)

            # MULTI-INSURED: Show insured header + subheader (Ramo + Nro Póliza)
            # Structure: SECTION START -> HEADER ASEGURADO -> SUBHEADER -> ...
            if has_multi_insured and insured_name:
                # Header Asegurado (18pt, spacing after 2pt)
                ih = _new_para()
                _set_style(ih, 'Heading')
                _noto_run(ih, insured_name, size_pt=18, bold=True)
                _set_spacing(ih, None, 2)

            # Process each policy block
            for policy_idx, (policy_id, policy_entries) in enumerate(policies_map.items()):
                # Ramo + Nro póliza info
                first_entry = policy_entries[0]['entry']
                first_branch = first_entry.get('branch', '')
                first_key = policy_entries[0]['key']
                first_numero = first_entry.get('numero', '')
                if not first_numero and first_key in policyNumbers:
                    first_numero = policyNumbers.get(first_key, '')

                poliza_label = f"PÓLIZA {first_branch.upper()}"
                if first_numero:
                    poliza_label += f" {first_numero}"

                # VARIOS ASEGURADOS: Subheader = 14pt, 12pt spacing After, Single line spacing
                # UN ASEGURADO: Header = 18pt, 12pt spacing After, Single line spacing
                sh = _new_para()
                _set_style(sh, 'Heading')
                if has_multi_insured:
                    header_size = 14
                else:
                    header_size = 18
                _noto_run(sh, poliza_label, size_pt=header_size, bold=True)
                _set_spacing(sh, None, 12)

                # Render all tables for this policy
                for table_idx, item in enumerate(policy_entries):
                    entry = item['entry']
                    entry_key = item['key']

                    branch = entry.get('branch', '')
                    numero = entry.get('numero', '')

                    if not numero and entry_key in policyNumbers:
                        numero = policyNumbers.get(entry_key, '')

                    current_mode = entry.get('mode', 'single')
                    current_is_individual = current_mode == 'by_endorsee'

                    has_groups = isinstance(entry.get('groups'), list) and len(entry.get('groups', [])) > 0
                    has_columns = isinstance(entry.get('columns'), list) and len(entry.get('columns', [])) > 0
                    has_rows = isinstance(entry.get('rows'), list) and len(entry.get('rows', [])) > 0

                    is_trec_native = branch.upper() == 'TREC' and current_mode != 'by_endorsee'

                    table_added = False
                    if is_trec_native and has_groups:
                        groups = entry.get('groups', [])
                        table_added = _add_table_trec(branch, numero, groups, currency, hasMultipleInsured=has_multi_insured)
                    elif current_is_individual and has_groups:
                        groups = entry.get('groups', [])
                        table_added = _add_table_by_endorsee(branch, numero, groups, currency, hasMultipleInsured=has_multi_insured)
                    elif has_columns and has_rows:
                        columns = entry.get('columns', [])
                        rows = entry.get('rows', [])
                        result = _add_table_dynamic(branch, numero, columns, rows, currency, hasMultipleInsured=has_multi_insured)
                        table_added = result[0] if isinstance(result, tuple) else result
                    else:
                        has_rows = isinstance(entry.get('rows'), list) and len(entry.get('rows', [])) > 0
                        if has_rows:
                            rows = entry.get('rows', [])
                            table_added = _add_table_3col(branch, numero, rows, currency, hasMultipleInsured=has_multi_insured)

                    # Insert ZeroWidthSpace (8pt) AFTER each table in Individual mode (except last)
                    # Spec: Spacing After 8pt, Line Spacing Single
                    if (table_added and current_is_individual and 
                        table_idx < len(policy_entries) - 1 and
                        len(policy_entries) > 1):
                        tables_in_body = body.findall(_qn('w:tbl'))
                        if tables_in_body:
                            _add_spacer(tables_in_body[-1], 16)

                # Add page break between policy blocks (except last policy)
                if policy_idx < len(policies_map) - 1:
                    _add_page_break()

        doc.save(savePath)
        return {'success': True, 'path': savePath, 'error': ''}

    except Exception as error:
        return {'success': False, 'path': '', 'error': str(error)}
